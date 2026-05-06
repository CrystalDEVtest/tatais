from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.db.models import Count, Avg, F, Q, DurationField
from django.db.models.functions import TruncDate, ExtractHour
from django.utils import timezone
from datetime import timedelta
import json
import csv
import io
from django.http import HttpResponse

from tickets.models import Ticket
from accounts.models import User

SERVICE_TYPE_LABELS = {
    'internet': 'Интернет',
    'tv': 'Телевидение',
    'phone': 'Телефония',
    'iptv': 'Интернет+ТВ',
    'internet_phone': 'Интернет+Телефон',
    'triple': 'Интернет+ТВ+Телефон',
    'other': 'Другое',
}


@login_required
def reports_dashboard(request):
    if request.user.role not in ('engineer', 'dispatcher', 'admin'):
        from django.contrib import messages
        messages.error(request, 'У вас нет доступа.')
        return redirect('home')
    
    days = int(request.GET.get('days', 30))
    start_date = timezone.now() - timedelta(days=days)
    
    total_tickets = Ticket.objects.filter(created_at__gte=start_date).count()
    completed_tickets = Ticket.objects.filter(status='completed', completed_at__gte=start_date).count()
    in_progress = Ticket.objects.filter(status__in=('assigned', 'in_progress', 'parts_needed')).count()
    new_tickets = Ticket.objects.filter(status='new').count()
    completion_rate = (completed_tickets / total_tickets * 100) if total_tickets > 0 else 0
    
    # Статусы — для графика
    status_choices_dict = dict(Ticket.STATUS_CHOICES)
    status_data_raw = dict(
        Ticket.objects.filter(created_at__gte=start_date)
        .values('status').annotate(count=Count('id')).values_list('status', 'count')
    )
    status_labels = json.dumps([status_choices_dict.get(k, k) for k in status_data_raw.keys()])
    status_values = json.dumps(list(status_data_raw.values()))
    
    # Типы услуг — для графика
    service_type_labels_map = {
        'internet': 'Интернет', 'tv': 'ТВ', 'phone': 'Телефония',
        'iptv': 'Интернет+ТВ', 'internet_phone': 'Интернет+Тел',
        'triple': 'Интернет+ТВ+Тел', 'other': 'Другое',
    }
    service_data_raw = dict(
        Ticket.objects.filter(created_at__gte=start_date)
        .values('service_type').annotate(count=Count('id')).values_list('service_type', 'count')
    )
    service_labels = json.dumps([service_type_labels_map.get(k, k) for k in service_data_raw.keys()])
    service_values = json.dumps(list(service_data_raw.values()))
    
    # Города — для графика
    city_data_raw = dict(
        Ticket.objects.filter(created_at__gte=start_date)
        .values('city').annotate(count=Count('id')).values_list('city', 'count')
    )
    city_labels = json.dumps(list(city_data_raw.keys()))
    city_values = json.dumps(list(city_data_raw.values()))
    
    # Нагрузка инженеров
    if request.user.role in ('dispatcher', 'admin'):
        engineer_stats = list(User.objects.filter(role='engineer').annotate(
            active_tickets=Count('assigned_tickets', filter=Q(assigned_tickets__status__in=('assigned', 'in_progress', 'parts_needed'))),
            completed_count=Count('assigned_tickets', filter=Q(assigned_tickets__status='completed', assigned_tickets__completed_at__gte=start_date))
        ).order_by('-active_tickets'))
        
        engineer_workload = []
        for eng in engineer_stats:
            load = min(eng.active_tickets * 10, 100)
            engineer_workload.append({
                'first_name': eng.first_name,
                'last_name': eng.last_name,
                'get_full_name': eng.full_name,
                'active_count': eng.active_tickets,
                'completed_count': eng.completed_count,
                'load_percent': load,
            })
    else:
        engineer_workload = None
        engineer_stats = None
    
    avg_completion = Ticket.objects.filter(
        status='completed', completed_at__isnull=False, created_at__gte=start_date
    ).annotate(duration=F('completed_at') - F('created_at')).aggregate(avg=Avg('duration'))['avg']
    
    recent_tickets = Ticket.objects.all().order_by('-created_at')[:20]
    
    rating_stats = Ticket.objects.filter(
        status='completed', customer_rating__isnull=False
    ).aggregate(avg_rating=Avg('customer_rating'), total_rated=Count('id'))
    
    # Динамика по дням
    daily_tickets = (
        Ticket.objects.filter(created_at__gte=start_date)
        .annotate(date=TruncDate('created_at'))
        .values('date').annotate(count=Count('id')).order_by('date')
    )
    daily_labels = json.dumps([str(d['date']) for d in daily_tickets])
    daily_values = json.dumps([d['count'] for d in daily_tickets])
    
    context = {
        'total_tickets': total_tickets,
        'completed_tickets': completed_tickets,
        'in_progress': in_progress,
        'in_progress_tickets': in_progress,
        'new_tickets': new_tickets,
        'completion_rate': round(completion_rate, 1),
        'status_data': status_data_raw,
        'status_labels': status_labels,
        'status_values': status_values,
        'service_data': service_data_raw,
        'service_labels': service_labels,
        'service_values': service_values,
        'city_data': city_data_raw,
        'city_labels': city_labels,
        'city_values': city_values,
        'engineer_stats': engineer_stats,
        'engineer_workload': engineer_workload,
        'avg_completion': avg_completion,
        'avg_rating': round(rating_stats['avg_rating'], 1) if rating_stats['avg_rating'] else '—',
        'recent_tickets': recent_tickets,
        'rating_stats': rating_stats,
        'daily_labels': daily_labels,
        'daily_values': daily_values,
        'days': days,
        'status_choices': dict(Ticket.STATUS_CHOICES),
        'service_type_labels': service_type_labels_map,
    }
    
    return render(request, 'reports/dashboard.html', context)


@login_required
def export_excel(request):
    if request.user.role not in ('engineer', 'dispatcher', 'admin'):
        return redirect('home')
    
    days = int(request.GET.get('days', 30))
    start_date = timezone.now() - timedelta(days=days)
    tickets = Ticket.objects.filter(created_at__gte=start_date).select_related('customer', 'assigned_engineer')
    
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Заявки'
    
    headers = ['№', 'Номер', 'Тип услуги', 'Статус', 'Приоритет', 'Город', 'Адрес', 'Заказчик', 'Инженер', 'Дата создания', 'Дата завершения', 'Оценка']
    ws.append(headers)
    
    for i, t in enumerate(tickets, 1):
        ws.append([
            i, t.ticket_number, SERVICE_TYPE_LABELS.get(t.service_type, t.service_type), t.get_status_display(),
            t.get_priority_display(), t.city, t.address,
            str(t.customer) if t.customer else (t.guest_name or 'Гость'),
            str(t.assigned_engineer) if t.assigned_engineer else 'Не назначен',
            t.created_at.strftime('%d.%m.%Y %H:%M') if t.created_at else '',
            t.completed_at.strftime('%d.%m.%Y %H:%M') if t.completed_at else '',
            t.customer_rating or '',
        ])
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="otchet_zayavki.xlsx"'
    wb.save(response)
    return response


@login_required
def export_csv(request):
    if request.user.role not in ('engineer', 'dispatcher', 'admin'):
        return redirect('home')
    
    days = int(request.GET.get('days', 30))
    start_date = timezone.now() - timedelta(days=days)
    tickets = Ticket.objects.filter(created_at__gte=start_date).select_related('customer', 'assigned_engineer')
    
    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="otchet_zayavki.csv"'
    
    writer = csv.writer(response, delimiter=';')
    writer.writerow(['№', 'Номер', 'Тип услуги', 'Статус', 'Приоритет', 'Город', 'Адрес', 'Заказчик', 'Инженер', 'Дата создания', 'Дата завершения', 'Оценка'])
    
    for i, t in enumerate(tickets, 1):
        writer.writerow([
            i, t.ticket_number, SERVICE_TYPE_LABELS.get(t.service_type, t.service_type), t.get_status_display(),
            t.get_priority_display(), t.city, t.address,
            str(t.customer) if t.customer else (t.guest_name or 'Гость'),
            str(t.assigned_engineer) if t.assigned_engineer else 'Не назначен',
            t.created_at.strftime('%d.%m.%Y %H:%M') if t.created_at else '',
            t.completed_at.strftime('%d.%m.%Y %H:%M') if t.completed_at else '',
            t.customer_rating or '',
        ])
    
    return response


@login_required
def export_word(request):
    if request.user.role not in ('engineer', 'dispatcher', 'admin'):
        return redirect('home')
    
    days = int(request.GET.get('days', 30))
    start_date = timezone.now() - timedelta(days=days)
    tickets = Ticket.objects.filter(created_at__gte=start_date).select_related('customer', 'assigned_engineer')
    
    from docx import Document
    doc = Document()
    doc.add_heading('Отчёт по заявкам — ТатАИСнефть', level=1)
    doc.add_paragraph(f'Период: {start_date.strftime("%d.%m.%Y")} — {timezone.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Всего заявок: {tickets.count()}')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(['Номер', 'Тип', 'Статус', 'Город', 'Заказчик', 'Инженер']):
        hdr_cells[i].text = h
    
    for t in tickets[:100]:
        row_cells = table.add_row().cells
        row_cells[0].text = t.ticket_number
        row_cells[1].text = SERVICE_TYPE_LABELS.get(t.service_type, t.service_type)
        row_cells[2].text = t.get_status_display()
        row_cells[3].text = t.city
        row_cells[4].text = str(t.customer) if t.customer else (t.guest_name or 'Гость')
        row_cells[5].text = str(t.assigned_engineer) if t.assigned_engineer else 'Не назначен'
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="otchet_zayavki.docx"'
    doc.save(response)
    return response